# import os
# from pdf2image import convert_from_path
# from paddleocr import PaddleOCR
# from typing import Union
# import requests
# import tempfile

# # # -----------------------------
# # # OCR INITIALIZATION
# # # -----------------------------
# # ocr = PaddleOCR(
# #     lang="en",
# #     use_textline_orientation=True
# # )

# # # -----------------------------
# # # PDF → TEXT
# # # -----------------------------
# # def extract_from_pdf(pdf_path: str) -> str:
# #     print('pdf_path=========', pdf_path)
# #     if not os.path.exists(pdf_path):
# #         raise FileNotFoundError(pdf_path)

# #     images = convert_from_path(pdf_path, dpi=300)
# #     extracted_text = []

# #     with tempfile.TemporaryDirectory() as tmpdir:
# #         for i, img in enumerate(images):
# #             img_path = os.path.join(tmpdir, f"page_{i}.png")
# #             img.save(img_path, "PNG")

# #             result = ocr.ocr(img_path)

# #             for page in result:
# #                 if isinstance(page, dict):
# #                     for t in page.get("rec_texts", []):
# #                         if t.strip():
# #                             extracted_text.append(t)

# #     return "\n".join(extracted_text)


# # =========================================================
# # OCR INITIALIZATION (PDF ONLY)
# # =========================================================
# ocr = PaddleOCR(lang="en", use_textline_orientation=True)


# # =========================================================
# # INPUT NORMALIZER
# # =========================================================
# def normalize_input(input_value: Union[str, bytes]) -> str:
#     """
#     Returns a local file path for any input:
#     - local path
#     - http / https / S3 URL
#     - raw bytes
#     """

#     # Local file
#     if isinstance(input_value, str) and os.path.exists(input_value):
#         return input_value

#     # URL
#     if isinstance(input_value, str) and input_value.startswith(("http://", "https://")):
#         r = requests.get(input_value, timeout=60)
#         r.raise_for_status()

#         suffix = os.path.splitext(input_value.split("?")[0])[1] or ".bin"
#         fd, path = tempfile.mkstemp(suffix=suffix)
#         with os.fdopen(fd, "wb") as f:
#             f.write(r.content)
#         return path

#     # Raw bytes
#     if isinstance(input_value, (bytes, bytearray)):
#         fd, path = tempfile.mkstemp(suffix=".bin")
#         with os.fdopen(fd, "wb") as f:
#             f.write(input_value)
#         return path

#     raise ValueError("Unsupported input type")


# # =========================================================
# # TEXT EXTRACTORS
# # =========================================================
# def extract_text_from_pdf(path: str) -> str:
#     images = convert_from_path(path, dpi=300)
#     text = []

#     with tempfile.TemporaryDirectory() as tmpdir:
#         for i, img in enumerate(images):
#             img_path = os.path.join(tmpdir, f"page_{i}.png")
#             img.save(img_path, "PNG")

#             result = ocr.ocr(img_path)
#             for page in result:
#                 if isinstance(page, dict):
#                     text.extend(t for t in page.get("rec_texts", []) if t.strip())

#     return "\n".join(text)


# def extract_text_from_docx(path: str) -> str:
#     doc = Document(path)
#     return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# def extract_text_from_txt(path: str) -> str:
#     with open(path, "r", encoding="utf-8", errors="ignore") as f:
#         return f.read()


# # =========================================================
# # SINGLE ENTRY POINT (AUTO-DETECT)
# # =========================================================
# def extract_raw_text(input_value: Union[str, bytes]) -> str:
#     print('input_value=========', input_value)  
#     path = normalize_input(input_value)
#     print('path=========', path)    
#     ext = os.path.splitext(path.lower())[1]
#     print('ext=========', ext)  
#     if ext == ".pdf":
#         return extract_text_from_pdf(path)

#     if ext in (".doc", ".docx"):
#         return extract_text_from_docx(path)

#     if ext == ".txt":
#         return extract_text_from_txt(path)

#     raise ValueError(f"Unsupported file type: {ext}")


import os
import tempfile
import mimetypes

def detect_file_type(filename: str, content_type: str) -> str:
    ext = os.path.splitext(filename.lower())[1]

    if ext in [".pdf"]:
        return "pdf"
    if ext in [".docx"]:
        return "docx"
    if ext in [".txt", ".csv"]:
        return "text"
    if ext in [".xlsx", ".xls"]:
        return "excel"
    if ext in [".jpg", ".jpeg", ".png"]:
        return "image"
    if ext in [".json"]:
        return "json"
    if ext in [".xml"]:
        return "xml"

    raise ValueError("Unsupported file type")

def save_temp_file(content: bytes, suffix: str) -> str:
    fd, path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "wb") as f:
        f.write(content)
    return path


import pdfplumber
from docx import Document
import pandas as pd
import json
import xml.etree.ElementTree as ET
from pdf2image import convert_from_path
from paddleocr import PaddleOCR

ocr = PaddleOCR(lang="en", use_textline_orientation=True)

def extract_pdf_text(path: str) -> str:
    # Try text-based first
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text.append(t)
    if text:
        return "\n".join(text)

    # Fallback to OCR
    images = convert_from_path(path, dpi=300)
    ocr_text = []
    for img in images:
        result = ocr.ocr(img)
        for page in result:
            if isinstance(page, dict):
                ocr_text.extend(page.get("rec_texts", []))
    return "\n".join(ocr_text)


def extract_docx_text(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_excel_text(path: str) -> str:
    dfs = pd.read_excel(path, sheet_name=None)
    text = []
    for name, df in dfs.items():
        text.append(df.to_string())
    return "\n".join(text)


def extract_json_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return json.dumps(json.load(f), indent=2)


def extract_xml_text(path: str) -> str:
    tree = ET.parse(path)
    return ET.tostring(tree.getroot(), encoding="unicode")


def extract_raw_text(filename: str, content_type: str, content: bytes) -> str:
    file_type = detect_file_type(filename, content_type)
    path = save_temp_file(content, suffix=os.path.splitext(filename)[1])

    if file_type == "pdf":
        return extract_pdf_text(path)
    if file_type == "docx":
        return extract_docx_text(path)
    if file_type == "text":
        return extract_text_file(path)
    if file_type == "excel":
        return extract_excel_text(path)
    if file_type == "json":
        return extract_json_text(path)
    if file_type == "xml":
        return extract_xml_text(path)
    if file_type == "image":
        result = ocr.ocr(path)
        text = []
        for page in result:
            if isinstance(page, dict):
                text.extend(page.get("rec_texts", []))
        return "\n".join(text)

    raise ValueError("Unhandled file type")

